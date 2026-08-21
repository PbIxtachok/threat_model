# -*- coding: utf-8 -*-
"""Сборка итогового DOCX «Модель угроз безопасности информации».

Структура повторяет шаблон (титульный лист, 7 разделов, 4 приложения).
Генерация полностью программная (python-docx), LLM на этом этапе не используется.
"""
from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from . import dictionaries as dicts
from .schema import Profile

FONT = "Times New Roman"


# ----------------------------------------------------------------------
def _setup_styles(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for name, size, bold in (("Heading 1", 16, True), ("Heading 2", 14, True), ("Heading 3", 12, True)):
        h = doc.styles[name]
        h.font.name = FONT
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = None
        h.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def _p(doc, text: str, *, bold=False, align=None, size=None, indent=True):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    elif indent:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.first_line_indent = Cm(1.25)
    run = par.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return par


def _table(doc, headers: List[str], rows: List[List[str]], widths: List[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for par in cells[i].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(10)
                    r.font.name = FONT
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def _page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _multiline(items: List[str]) -> str:
    return "\n".join(f"– {i}" for i in items) if items else "—"


def _llm_paragraphs(doc, text: str) -> None:
    """Добавляет непустые абзацы LLM-текста (разделитель абзацев — пустая строка)."""
    for para in (text or "").split("\n\n"):
        para = para.strip()
        if para:
            _p(doc, para)


# ======================================================================
def build_document(profile: Profile, sections: Dict[str, Dict[str, Any]], out_path: str) -> str:
    p = profile
    doc = Document()
    _setup_styles(doc)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(2.5), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

    # ------------------------------------------------ титульный лист
    _p(doc, "УТВЕРЖДАЮ", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, p.approver_position or "___________________________", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, "____________ /____________/", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, f"«___» __________ {p.year or '20__'} г.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    for _ in range(6):
        doc.add_paragraph()
    _p(doc, "МОДЕЛЬ УГРОЗ БЕЗОПАСНОСТИ ИНФОРМАЦИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    _p(doc, p.object_name or "Наименование информационной системы", bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    if p.operator_name:
        _p(doc, f"Оператор: {p.operator_name}", align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        doc.add_paragraph()
    _p(doc, f"{p.city or ''} {p.year or ''}".strip(), align=WD_ALIGN_PARAGRAPH.CENTER)
    _page_break(doc)

    # ------------------------------------------------ 1. Общие положения
    s1 = sections.get("s1_general", {})
    doc.add_heading("1. Общие положения", level=1)
    _p(doc, f"Настоящая Модель угроз безопасности информации (далее – Модель угроз) "
            f"разработана в отношении {p.object_name} (далее – Система) и содержит "
            f"результаты оценки угроз безопасности информации, определение возможных "
            f"негативных последствий, объектов воздействия, источников угроз, способов "
            f"их реализации и актуальных угроз безопасности информации.")
    _llm_paragraphs(doc, s1.get("intro_text", ""))
    _p(doc, "Модель угроз разработана в соответствии со следующими нормативными правовыми "
            "актами и методическими документами:")
    norms = [
        "Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных»;",
        "постановление Правительства Российской Федерации от 01.11.2012 № 1119 "
        "«Об утверждении требований к защите персональных данных при их обработке "
        "в информационных системах персональных данных»;",
        "приказ ФСТЭК России от 18.02.2013 № 21 «Об утверждении Состава и содержания "
        "организационных и технических мер по обеспечению безопасности персональных данных "
        "при их обработке в информационных системах персональных данных»;",
    ]
    if p.has_crypto:
        norms.append(
            "приказ ФСБ России от 10.07.2014 № 378 «Об утверждении Состава и "
            "содержания организационных и технических мер по обеспечению безопасности "
            "персональных данных при их обработке в информационных системах "
            "персональных данных с использованием средств криптографической защиты "
            "информации, необходимых для выполнения установленных Правительством "
            "Российской Федерации требований к защите персональных данных для "
            "каждого из уровней защищенности»;")
    norms += [
        "Методический документ «Методика оценки угроз безопасности информации» "
        "(утв. ФСТЭК России 05.02.2021);",
        "Банк данных угроз безопасности информации ФСТЭК России (bdu.fstec.ru).",
    ]
    for norm in norms:
        _p(doc, norm)
    _p(doc, f"Оператор Системы: {p.operator_name or '—'}.")
    if p.responsible:
        _p(doc, f"Подразделение (лицо), ответственное за обеспечение безопасности информации: {p.responsible}.")
    if p.developer_org:
        _p(doc, f"Организация, выполнившая разработку Модели угроз: {p.developer_org}.")
    _p(doc, "Модель угроз подлежит пересмотру при изменении архитектуры Системы, состава "
            "обрабатываемой информации, появлении новых угроз безопасности информации, "
            "а также по результатам контроля состояния защищённости.")
    _page_break(doc)

    # ------------------------------------------------ 2. Описание системы
    s2 = sections.get("s2_description", {})
    doc.add_heading("2. Описание информационной системы", level=1)
    doc.add_heading("2.1. Наименование, назначение и классификация", level=2)
    _p(doc, f"Наименование: {p.object_name}.")
    _p(doc, f"Тип системы: {p.system_type or '—'}.")
    if p.protection_level:
        basis = f" ({p.classification_basis})" if p.classification_basis else ""
        _p(doc, f"Установленный уровень защищённости (класс): {p.protection_level}{basis}.")
    if p.purpose:
        _p(doc, f"Назначение Системы: {p.purpose}")
    if p.business_processes:
        _p(doc, f"Основные процессы, обеспечиваемые Системой: {p.business_processes}")
    if p.scale:
        _p(doc, f"Масштаб и размещение: {p.scale}")
    _llm_paragraphs(doc, s2.get("purpose_text", ""))

    doc.add_heading("2.2. Обрабатываемая информация", level=2)
    _p(doc, "В Системе обрабатываются следующие виды информации:")
    for k in p.info_kinds:
        _p(doc, f"– {k};")
    if p.pdn_categories:
        _p(doc, f"Категории персональных данных: {', '.join(p.pdn_categories)}.")
        if p.pdn_subjects:
            _p(doc, f"Субъекты персональных данных: {', '.join(p.pdn_subjects)}.")
        if p.pdn_volume:
            _p(doc, f"Объём обрабатываемых персональных данных: {p.pdn_volume.lower()}.")
    _llm_paragraphs(doc, s2.get("info_text", ""))

    doc.add_heading("2.3. Архитектура и условия функционирования", level=2)
    _llm_paragraphs(doc, s2.get("architecture_text", ""))
    _p(doc, "Состав основных компонентов Системы приведён в Приложении 2.")
    _p(doc, f"Используемые интерфейсы и каналы взаимодействия: "
            f"{', '.join(p.network_interfaces) or '—'}.")
    flags = [
        ("подключение к сети «Интернет»", p.has_internet),
        ("беспроводные сети", p.has_wireless),
        ("удалённый доступ", p.has_remote_access),
        ("привлечение подрядных организаций", p.has_contractors),
        ("информационное взаимодействие с внешними системами", p.has_external_integrations),
    ]
    used = [n for n, v in flags if v]
    not_used = [n for n, v in flags if not v]
    if used:
        _p(doc, "В Системе используется: " + "; ".join(used) + ".")
    if not_used:
        _p(doc, "Не используется: " + "; ".join(not_used) + ".")
    if p.cloud_model and p.cloud_model != "Не используется":
        _p(doc, f"Модель размещения: {p.cloud_model}. {p.cloud_details}".strip())
    if p.security_tools:
        _p(doc, "Применяемые средства защиты информации:")
        for tool in p.security_tools:
            _p(doc, f"– {tool};")
        if p.has_crypto:
            _p(doc, "В Системе применяются средства криптографической защиты "
                    "информации (СКЗИ).")

    doc.add_heading("2.4. Пользователи системы", level=2)
    _p(doc, "К работе с Системой допущены следующие группы пользователей:")
    for g in p.user_groups:
        _p(doc, f"– {g};")
    if p.users_notes:
        _p(doc, p.users_notes)
    _llm_paragraphs(doc, s2.get("users_text", ""))
    _page_break(doc)

    # ------------------------------------------------ 3. Негативные последствия
    s3 = sections.get("s3_consequences", {})
    doc.add_heading("3. Возможные негативные последствия от реализации угроз", level=1)
    _p(doc, "По результатам анализа исходных данных определены следующие виды риска "
            "(ущерба) и типовые негативные последствия, актуальные для Системы "
            "(в соответствии с таблицей 4.1 Методики оценки угроз безопасности информации):")
    rows = []
    for r in s3.get("rows", []):
        rows.append([r["id"], r["name"], _multiline(r["consequences"])])
    custom = s3.get("custom", [])
    if custom:
        rows.append(["—", "Дополнительно определённые оператором", _multiline(custom)])
    _table(doc, ["Вид риска", "Наименование", "Негативные последствия"], rows, widths=[2.2, 5.5, 9.5])
    rationale = (s3.get("rationale") or "").strip()
    if rationale:
        doc.add_paragraph()
        _llm_paragraphs(doc, rationale)
    _page_break(doc)

    # ------------------------------------------------ 4. Объекты воздействия
    s4 = sections.get("s4_impact_objects", {})
    doc.add_heading("4. Объекты воздействия угроз безопасности информации", level=1)
    _p(doc, "На основе анализа архитектуры Системы и обрабатываемой информации определены "
            "объекты воздействия и виды воздействия на них:")
    _llm_paragraphs(doc, s4.get("analysis_text", ""))
    rows = [[r["component"], r["ctype"], r["location"], _multiline(r["impacts"])]
            for r in s4.get("rows", [])]
    _table(doc, ["Объект воздействия", "Тип", "Размещение", "Виды воздействия"], rows,
           widths=[4.0, 3.5, 3.0, 6.7])
    _page_break(doc)

    # ------------------------------------------------ 5. Источники угроз
    s5 = sections.get("s5_intruders", {})
    doc.add_heading("5. Источники угроз безопасности информации", level=1)
    doc.add_heading("5.1. Антропогенные источники (нарушители)", level=2)
    _p(doc, "По результатам оценки определены следующие актуальные нарушители:")
    rows = []
    for t in s5.get("intruders", []):
        rows.append([
            f"{t['id']}. {t['name']}",
            t["category"],
            t["level"],
            _multiline(t.get("goals", [])),
            ", ".join(t.get("relevant_damages", [])) or "—",
        ])
    _table(doc, ["Вид нарушителя", "Категория", "Уровень", "Возможные цели", "Виды ущерба"],
           rows, widths=[4.5, 2.2, 1.6, 6.4, 2.0])
    _llm_paragraphs(doc, s5.get("intruders_text", ""))
    ml = s5.get("max_level", {})
    if ml:
        doc.add_paragraph()
        _p(doc, f"Максимальный уровень возможностей актуальных нарушителей — {ml.get('id')} "
                f"({ml.get('title', '').lower()}).")
        if ml.get("description"):
            _p(doc, ml["description"])
    coll = s5.get("collusion", [])
    if coll:
        _p(doc, "В соответствии с Методикой учтены следующие возможности сговора нарушителей:")
        for c in coll:
            _p(doc, f"– вид {c['a_id']} ({c['a_name']}) — с видом {c['b_id']} ({c['b_name']});")
    if s5.get("excluded_reason"):
        _p(doc, f"Обоснование исключения отдельных видов нарушителей: {s5['excluded_reason']}")
    doc.add_heading("5.2. Техногенные источники", level=2)
    _p(doc, "К техногенным источникам угроз относятся отказы и сбои программных и "
            "программно-аппаратных средств, средств обеспечения функционирования, "
            "а также ошибки, возникающие в процессе эксплуатации. Угрозы, обусловленные "
            "техногенными источниками, рассматриваются в рамках обеспечения отказоустойчивости "
            "и резервирования и не являются предметом настоящей Модели угроз.")
    _page_break(doc)

    # ------------------------------------------------ 6. Способы реализации
    s6 = sections.get("s6_ways", {})
    doc.add_heading("6. Способы реализации (возникновения) угроз", level=1)
    _p(doc, "Основными способами реализации угроз безопасности информации являются:")
    for m in s6.get("modules", []):
        _p(doc, f"– {m};")
    _p(doc, "Актуальные способы реализации угроз определяются с учётом имеющихся "
            "интерфейсов объектов воздействия: " + (", ".join(s6.get("interfaces", [])) or "—") + ".")
    _llm_paragraphs(doc, s6.get("ways_text", ""))
    _p(doc, "Детализация способов (техник и тактик) реализации применительно к актуальным "
            "угрозам приведена в разделе 7.2 и Приложении 3.")
    _page_break(doc)

    # ------------------------------------------------ 7. Актуальные угрозы
    s7 = sections.get("s7_ubi", {})
    s8 = sections.get("s8_scenarios", {})
    doc.add_heading("7. Актуальные угрозы безопасности информации", level=1)
    doc.add_heading("7.1. Перечень актуальных угроз", level=2)
    actual = s7.get("all", [])
    actual = [i for i in actual if i.get("matches")]
    _p(doc, f"По результатам оценки возможности реализации угроз из Банка данных угроз "
            f"безопасности информации ФСТЭК России актуальными для Системы признаны "
            f"{len(actual)} угроз:")
    _llm_paragraphs(doc, s7.get("intro_text", ""))
    rows = [[i.get("ubi_code", i.get("Number")), i.get("text", ""),
             i.get("confidence", ""), i.get("explanation", "")] for i in actual]
    _table(doc, ["УБИ", "Наименование угрозы", "Уверенность", "Обоснование актуальности"],
           rows, widths=[2.2, 6.0, 2.3, 6.7])
    _llm_paragraphs(doc, s7.get("summary_text", ""))

    doc.add_heading("7.2. Сценарии реализации актуальных угроз", level=2)
    pairs = s8.get("pairs", [])
    if pairs:
        _p(doc, "Для актуальных угроз определены возможные техники (способы) их реализации "
                "(полная матрица приведена в Приложении 4):")
        scenario_texts = s8.get("scenario_texts", {}) or {}
        by_threat: Dict[str, List[Dict[str, Any]]] = {}
        for pr in pairs:
            by_threat.setdefault(pr["threat_number"], []).append(pr)
        for num in sorted(by_threat, key=lambda x: int(x) if str(x).isdigit() else 0):
            items = by_threat[num]
            doc.add_heading(f"УБИ.{int(num):03d}. {items[0]['threat_text']}", level=3)
            _llm_paragraphs(doc, scenario_texts.get(str(num), ""))
            rows = [[i["tactic_id"], i["tactic_category"], i["tactic"][:200], i["explanation"]]
                    for i in items]
            _table(doc, ["Техника", "Тактика", "Описание", "Пояснение"], rows,
                   widths=[1.8, 3.5, 6.0, 5.9])
    else:
        _p(doc, s8.get("note", "Сценарии не сформированы."))

    # --- 7.3. Тип актуальных угроз и требуемый класс СКЗИ (приказ ФСБ № 378) ---
    skzi = s7.get("skzi")
    if skzi:
        doc.add_heading("7.3. Тип актуальных угроз и требуемый класс СКЗИ", level=2)
        _p(doc, "В Системе для защиты персональных данных применяются средства "
                "криптографической защиты информации. В соответствии с постановлением "
                "Правительства Российской Федерации от 01.11.2012 № 1119 для уровня "
                f"защищённости {skzi.get('uz', '')} актуальными являются угрозы "
                "следующих типов: 1-й тип — угрозы, связанные с наличием "
                "недекларированных возможностей (НДВ) в системном ПО; 2-й тип — "
                "с НДВ в прикладном ПО; 3-й тип — не связанные с НДВ. Требуемый "
                "класс средств криптографической защиты информации определён в "
                "соответствии с приказом ФСБ России от 10.07.2014 № 378:")
        rows = [[f"{r['type']}-й тип", r["class"]] for r in skzi.get("rows", [])]
        _table(doc, ["Тип актуальных угроз", "Требуемый класс СКЗИ"], rows,
               widths=[8.6, 8.6])
        _p(doc, f"Минимально требуемый класс СКЗИ для Системы: {skzi.get('min_class', '—')}.")
        _llm_paragraphs(doc, (s7.get("skzi_text") or "").strip())
    _page_break(doc)

    # ------------------------------------------------ Приложение 1: сокращения
    doc.add_heading("Приложение 1. Перечень сокращений", level=1)
    abbr = dicts.misc()["abbreviations"]
    _table(doc, ["Сокращение", "Расшифровка"], [[a, b] for a, b in abbr], widths=[3.5, 13.7])
    _page_break(doc)

    # ------------------------------------------------ Приложение 2: компоненты
    doc.add_heading("Приложение 2. Состав компонентов информационной системы", level=1)
    rows = [[c.name, c.ctype, c.purpose or "—", c.location or "—"] for c in p.components]
    _table(doc, ["Наименование", "Тип", "Назначение", "Размещение"], rows,
           widths=[4.5, 3.7, 5.5, 3.5])
    _page_break(doc)

    # ------------------------------------------------ Приложение 3: угрозы × нарушители
    doc.add_heading("Приложение 3. Соотнесение актуальных угроз и нарушителей", level=1)
    intr = sections.get("s5_intruders", {}).get("intruders", [])
    intr_names = "; ".join(f"{t['id']} — {t['name']} ({t['level']})" for t in intr) or "—"
    _p(doc, f"Актуальные нарушители: {intr_names}.")
    max_lvl = sections.get("s5_intruders", {}).get("max_level", {}).get("id", "—")
    intr_ids = ", ".join(str(t["id"]) for t in intr if t.get("actual")) or "—"
    rows = [[i.get("ubi_code", ""), i.get("text", ""), max_lvl, f"Виды: {intr_ids}"]
            for i in actual]
    _table(doc, ["УБИ", "Угроза", "Уровень нарушителя", "Возможные нарушители"], rows,
           widths=[2.2, 8.0, 2.5, 4.5])
    _page_break(doc)

    # ------------------------------------------------ Приложение 4: матрица актуальности
    doc.add_heading("Приложение 4. Результаты оценки возможности реализации угроз", level=1)
    all_items = s7.get("all", [])
    rows = [[i.get("ubi_code", i.get("Number")),
             i.get("text", "")[:150],
             "Актуальна" if i.get("matches") else "Неактуальна",
             i.get("confidence", ""),
             i.get("explanation", "")[:300]] for i in all_items]
    _table(doc, ["УБИ", "Наименование", "Вывод", "Уверенность", "Обоснование"], rows,
           widths=[2.0, 5.5, 2.4, 2.2, 5.1])
    _page_break(doc)

    # ------------------------------------------------ Приложение 5: перечень актуальных УБИ
    doc.add_heading("Приложение 5. Перечень актуальных угроз безопасности информации", level=1)
    def _ubi_sort_key(item: Dict[str, Any]) -> int:
        num = str(item.get("Number", ""))
        return int(num) if num.isdigit() else 0
    actual_sorted = sorted(actual, key=_ubi_sort_key)
    if actual_sorted:
        rows = [[i.get("ubi_code", i.get("Number")), i.get("text", "")]
                for i in actual_sorted]
        _table(doc, ["Идентификатор УБИ", "Наименование УБИ"], rows,
               widths=[3.5, 13.7])
    else:
        _p(doc, "Актуальные угрозы не выявлены.")

    doc.save(out_path)
    return out_path
