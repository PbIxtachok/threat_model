# -*- coding: utf-8 -*-
"""Сквозной тест конвейера на mock-провайдере: профиль → все этапы → DOCX.

Запуск:  python -m pytest tests/ -v   (или  python tests/test_pipeline.py)
"""
from __future__ import annotations

import json
import sys
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from tm_core import storage
from tm_core.analysis import _norm_ubi_number
from tm_core.pipeline import start_or_resume
from tm_core.providers import PROVIDERS, MockProvider, safe_parse_json
from tm_core.schema import Component, Profile


class FlakyMockProvider(MockProvider):
    """Первые 2 JSON-вызова по УБИ возвращают мусор, далее — как MockProvider."""

    name = "mock-flaky"

    def __init__(self, cfg):
        super().__init__(cfg)
        self.ubi_json_calls = 0

    def generate(self, system: str, user: str, json_mode: bool = False) -> str:
        if json_mode and "Список УБИ" in user:
            self.ubi_json_calls += 1
            if self.ubi_json_calls <= 2:
                return "это не JSON {{{ [обрезано"
        return super().generate(system, user, json_mode=json_mode)


class GarbageMockProvider(MockProvider):
    """Всегда возвращает мусор на JSON-запросы (эмуляция неработающей LLM)."""

    name = "mock-garbage"

    def generate(self, system: str, user: str, json_mode: bool = False) -> str:
        if json_mode:
            return "это не JSON {{{ [обрезано"
        return super().generate(system, user, json_mode=json_mode)


class CountingMockProvider(MockProvider):
    """Считает JSON-вызовы по тактикам (проверка переобработки s8)."""

    name = "mock-counting"

    def __init__(self, cfg):
        super().__init__(cfg)
        self.tactic_json_calls = 0

    def generate(self, system: str, user: str, json_mode: bool = False) -> str:
        if json_mode and "tactic_id" in user:
            self.tactic_json_calls += 1
        return super().generate(system, user, json_mode=json_mode)


class SplitMockProvider(MockProvider):
    """Считает вызовы по имени модели и типу (json/text) — для теста
    двух конфигураций LLM (analysis vs text)."""

    name = "mock-split"
    calls = {}  # model -> {"json": n, "text": n}

    def generate(self, system: str, user: str, json_mode: bool = False) -> str:
        rec = self.calls.setdefault(self.cfg.model, {"json": 0, "text": 0})
        rec["json" if json_mode else "text"] += 1
        return super().generate(system, user, json_mode=json_mode)


PROVIDERS.setdefault("flaky", FlakyMockProvider)
PROVIDERS.setdefault("garbage", GarbageMockProvider)
PROVIDERS.setdefault("counting", CountingMockProvider)
PROVIDERS.setdefault("splitmock", SplitMockProvider)


def make_test_profile() -> Profile:
    return Profile(
        object_name="ИСПДн «Кадры и зарплата»",
        operator_name="ООО «Ромашка»",
        responsible="Отдел информационной безопасности",
        approver_position="Генеральный директор ООО «Ромашка»",
        city="г. Оренбург",
        year="2026",
        system_type="ИСПДн",
        protection_level="УЗ-3",
        purpose="Автоматизация кадрового учёта и расчёта заработной платы работников.",
        business_processes="Кадровый учёт; расчёт заработной платы; отчётность в госорганы.",
        scale="Одна площадка: серверная и АРМ в офисе оператора.",
        info_kinds=["Персональные данные", "Учетные данные (логины, пароли, ключи)"],
        pdn_categories=["Иные категории ПДн"],
        pdn_subjects=["Сотрудники оператора"],
        pdn_volume="Менее 100 000 субъектов",
        components=[
            Component("Сервер 1С", "Сервер приложений", "1С:ЗУП", "Серверная"),
            Component("СУБД PostgreSQL", "Сервер баз данных", "Хранение БД", "Серверная"),
            Component("АРМ бухгалтера", "АРМ пользователя", "Работа в 1С", "Офис"),
            Component("Межсетевой экран", "Межсетевой экран", "Защита периметра", "Серверная"),
        ],
        network_interfaces=["ЛВС", "Интернет"],
        has_internet=True,
        has_remote_access=False,
        has_contractors=True,
        user_groups=["Пользователи (операторы) системы", "Администраторы системы"],
        damage_types=["У1", "У2"],
        consequences=[
            "Нарушение конфиденциальности (утечка) персональных данных",
            "Нарушение законодательства Российской Федерации",
        ],
        intruder_ids=[3, 4, 9, 11, 12, 13],
        security_tools=["Антивирус", "Межсетевой экран", "СЗИ от НСД"],
    )


def test_full_pipeline():
    profile = make_test_profile()
    errors = profile.validate()
    assert not errors, f"Профиль должен быть валиден: {errors}"

    state = storage.new_job(profile, "тестовый профиль")
    llm_cfg = {"provider": "mock", "model": "mock", "max_tokens": 512}

    job = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final = job.run()

    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"
    assert final.docx_path and Path(final.docx_path).exists(), "DOCX не создан"
    assert Path(final.docx_path).stat().st_size > 20_000, "DOCX подозрительно мал"

    # проверка чекпоинтов
    for stage in ("s1_general", "s7_ubi", "s9_docx"):
        assert storage.load_stage_result(state.job_id, stage) is not None, f"Нет результата {stage}"

    s7 = storage.load_stage_result(state.job_id, "s7_ubi")
    assert s7["all"], "Нет результатов по УБИ"
    print(f"OK: {final.docx_path}; УБИ рассмотрено {len(s7['all'])}, актуальных {len(s7['actual'])}")
    return final


def test_resume():
    """Проверка возобновления: у существующего задания сбрасываем последний этап и дорабатываем."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест возобновления")
    llm_cfg = {"provider": "mock", "model": "mock"}

    # первый запуск — сразу останавливаем после старта s7 (эмулируем: выполняем всё, затем "ломаем" state)
    job = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final = job.run()
    assert final.status == "done"

    # эмулируем прерывание: помечаем s9 как невыполненный и убираем DOCX
    st = storage.load_job_state(state.job_id)
    st.stages["s9_docx"] = "pending"
    st.status = "paused"
    storage.save_job_state(st)

    job2 = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final2 = job2.run()
    assert final2.status == "done", f"Возобновление не сработало: {final2.error}"
    assert Path(final2.docx_path).exists()
    print(f"OK resume: {final2.docx_path}")
    return final2


def test_safe_parse_json_repairs():
    """Ремонт типовых дефектов JSON от LLM: обрезанный массив, trailing commas."""
    truncated = ('[{"Number": "1", "matches": true}, {"Number": "2", "matches": false}, '
                 '{"Number": "3", "mat')
    parsed = safe_parse_json(truncated)
    assert isinstance(parsed, list), "Обрезанный массив должен быть восстановлен"
    assert [i["Number"] for i in parsed] == ["1", "2"]

    commas = '[{"Number": "1", "matches": true,}, {"Number": "2",},]'
    parsed = safe_parse_json(commas)
    assert isinstance(parsed, list) and len(parsed) == 2, \
        "Trailing commas должны удаляться"

    assert safe_parse_json('{"a": 1}') == {"a": 1}
    assert safe_parse_json("совсем не json") is None


def test_norm_ubi_number():
    """Нормализация номеров УБИ: префиксы и ведущие нули отбрасываются."""
    assert _norm_ubi_number("УБИ.042") == "42"
    assert _norm_ubi_number("042") == "42"
    assert _norm_ubi_number("УБИ 42") == "42"
    assert _norm_ubi_number("") == ""


def test_retry_recovers():
    """Flaky-провайдер (мусор в первых 2 вызовах): ретраи спасают батч,
    в итоге нет записей «не удалось получить»."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест ретраев")
    llm_cfg = {"provider": "flaky", "model": "mock-flaky"}

    job = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"

    s7 = storage.load_stage_result(state.job_id, "s7_ubi")
    bad = [i for i in s7["all"] if "не удалось получить" in i.get("explanation", "").lower()]
    assert not bad, f"Остались необработанные УБИ: {[i['Number'] for i in bad]}"
    print(f"OK retry: УБИ рассмотрено {len(s7['all'])}")
    return final


def test_total_failure_marked():
    """Провайдер-мусор: этап завершается, записи помечаются llm_error
    с новым текстом объяснения, DOCX строится."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест полного отказа LLM")
    llm_cfg = {"provider": "garbage", "model": "mock-garbage"}

    job = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"
    assert final.docx_path and Path(final.docx_path).exists(), "DOCX не создан"

    s7 = storage.load_stage_result(state.job_id, "s7_ubi")
    llm_items = [i for i in s7["all"] if i.get("source") not in ("rules",)]
    assert llm_items, "Должны быть записи, обрабатывавшиеся через LLM"
    for i in llm_items:
        assert i.get("source") == "llm_error", f"УБИ {i['Number']}: source={i.get('source')}"
        assert "Не удалось получить корректный ответ LLM" in i.get("explanation", "")
        assert i.get("matches") is False and i.get("confidence") == "низкая"
    print(f"OK total failure: llm_error {len(llm_items)} из {len(s7['all'])}")
    return final


def test_resume_reprocesses_failed():
    """После неудачного прогона (мусор от LLM) смена провайдера на нормальный
    mock + возобновление: записи llm_error переобрабатываются."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест переобработки при возобновлении")

    # первый прогон — LLM возвращает мусор
    job = start_or_resume(state.job_id, {"provider": "garbage", "model": "mock-garbage-resume"},
                          ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"
    s7 = storage.load_stage_result(state.job_id, "s7_ubi")
    assert any(i.get("source") == "llm_error" for i in s7["all"]), \
        "Первый прогон должен содержать записи llm_error"

    # эмулируем возобновление: s7 сброшен, partial-файл оставлен
    st = storage.load_job_state(state.job_id)
    st.stages["s7_ubi"] = "pending"
    st.status = "paused"
    storage.save_job_state(st)
    Path(storage.job_dir(state.job_id) / "s7_ubi.json").unlink(missing_ok=True)

    job2 = start_or_resume(state.job_id, {"provider": "mock", "model": "mock"},
                           ubi_full_scan=False)
    final2 = job2.run()
    assert final2.status == "done", f"Возобновление не сработало: {final2.error}"
    s7b = storage.load_stage_result(state.job_id, "s7_ubi")
    bad = [i for i in s7b["all"] if i.get("source") in ("llm_missing", "llm_error")]
    assert not bad, f"Записи с ошибкой LLM не переобработаны: {[i['Number'] for i in bad]}"
    print(f"OK resume reprocess: УБИ {len(s7b['all'])}, актуальных {len(s7b['actual'])}")
    return final2


def test_s8_old_format_partials_reprocessed():
    """Partial-записи s8 СТАРОГО формата (без поля source, неудача выглядела как
    explanation="Нет ответа LLM") при возобновлении переобрабатываются."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест старого формата s8")

    job = start_or_resume(state.job_id, {"provider": "mock", "model": "mock-s8old"},
                          ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"
    s8 = storage.load_stage_result(state.job_id, "s8_scenarios")
    assert s8["pairs"], "Для теста нужны совпавшие пары УБИ-тактика"

    # эмулируем partial-файл от baseline-версии: у записей нет поля source,
    # неудача выглядела как explanation == "Нет ответа LLM"
    ppath = storage.partial_path(state.job_id, "s8_scenarios")
    old_items = storage.load_partials(state.job_id, "s8_scenarios")
    assert old_items, "Должны быть partial-записи s8"
    with open(ppath, "w", encoding="utf-8") as f:
        for i in old_items:
            i.pop("source", None)
            i["explanation"] = "Нет ответа LLM"
            f.write(json.dumps(i, ensure_ascii=False) + "\n")

    # эмулируем возобновление: s8/s9 сброшены, старые результаты удалены
    st = storage.load_job_state(state.job_id)
    st.stages["s8_scenarios"] = "pending"
    st.stages["s9_docx"] = "pending"
    st.status = "paused"
    storage.save_job_state(st)
    Path(storage.job_dir(state.job_id) / "s8_scenarios.json").unlink(missing_ok=True)
    Path(storage.job_dir(state.job_id) / "s9_docx.json").unlink(missing_ok=True)

    # уникальное имя модели — ключи кэша LLM не пересекаются между прогонами,
    # поэтому переобработка гарантированно идёт через provider.generate
    job2 = start_or_resume(state.job_id,
                           {"provider": "counting",
                            "model": f"mock-counting-{uuid.uuid4().hex[:8]}"},
                           ubi_full_scan=False)
    final2 = job2.run()
    assert final2.status == "done", f"Возобновление не сработало: {final2.error}"
    assert job2.provider is not None
    assert job2.provider.tactic_json_calls > 0, \
        "Записи старого формата (без source) должны переобрабатываться — " \
        "ожидались обращения к LLM"

    s8b = storage.load_stage_result(state.job_id, "s8_scenarios")
    assert s8b["pairs"], "Пары УБИ-тактика должны быть пересчитаны"
    expls = {str(i.get("explanation")) for i in s8b["pairs"]}
    assert "Нет ответа LLM" not in expls, "Старые записи не переобработаны"
    new_items = storage.load_partials(state.job_id, "s8_scenarios")
    assert any(i.get("source") == "llm" for i in new_items), \
        "В partial-файле должны появиться записи нового формата (source=llm)"
    print(f"OK old-format s8: вызовов LLM по тактикам {job2.provider.tactic_json_calls}, "
          f"пар {len(s8b['pairs'])}")
    return final2


def test_downstream_invalidation():
    """Повторное выполнение s7 инвалидирует downstream-этапы: s8/s9 переводятся
    в pending, их результаты удаляются и пересчитываются, DOCX собирается заново."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест инвалидации downstream")
    llm_cfg = {"provider": "mock", "model": "mock-invalidate"}

    job = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"

    job_dir = storage.job_dir(state.job_id)
    s8_orig = storage.load_stage_result(state.job_id, "s8_scenarios")
    docx_path = Path(final.docx_path)
    assert docx_path.exists()

    # помечаем s8 «битым» результатом и удаляем DOCX/s9: если downstream не
    # инвалидируется при повторном s7, эти артефакты останутся как есть
    storage.save_stage_result(state.job_id, "s8_scenarios", {"sentinel": True, "pairs": []})
    docx_path.unlink()
    Path(job_dir / "s9_docx.json").unlink(missing_ok=True)

    # пользователь сбрасывает s7 и возобновляет генерацию
    st = storage.load_job_state(state.job_id)
    st.stages["s7_ubi"] = "pending"
    st.status = "paused"
    storage.save_job_state(st)

    job2 = start_or_resume(state.job_id, llm_cfg, ubi_full_scan=False)
    final2 = job2.run()
    assert final2.status == "done", f"Возобновление не сработало: {final2.error}"

    s8_new = storage.load_stage_result(state.job_id, "s8_scenarios")
    assert s8_new is not None and "sentinel" not in s8_new, \
        "s8 должен быть пересчитан после повторного выполнения s7"
    assert s8_new.get("pairs") == s8_orig.get("pairs"), \
        "s8 пересчитан из сохранённых partial-чекпоинтов"
    assert storage.load_stage_result(state.job_id, "s9_docx") is not None, \
        "s9_docx должен быть пересчитан"
    assert Path(final2.docx_path).exists(), "DOCX должен быть собран заново"
    print(f"OK downstream invalidation: s8 пересчитан ({len(s8_new.get('pairs', []))} пар), "
          f"DOCX собран заново")
    return final2


def test_draft_roundtrip():
    """Черновик профиля: save_draft → load_draft (roundtrip)."""
    data = make_test_profile().to_dict()
    storage.save_draft(data)
    assert storage.DRAFT_FILE.exists(), "Файл черновика не создан"
    loaded = storage.load_draft()
    assert loaded == data, "Черновик должен читаться без потерь"
    print(f"OK draft roundtrip: {storage.DRAFT_FILE.name}")


def test_llm_settings_backward_compat():
    """Старый плоский формат llm_settings.json (ключ "provider") читается
    как единая конфигурация для обеих ролей; новый формат — как есть."""
    flat = {"provider": "mock", "model": "mock-flat", "max_tokens": 1024}
    with open(storage.SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(flat, f, ensure_ascii=False)
    loaded = storage.load_llm_settings()
    assert set(loaded) == {"analysis", "text"}, \
        f"Плоский dict должен разворачиваться в analysis+text: {loaded}"
    assert loaded["analysis"]["model"] == "mock-flat"
    assert loaded["text"]["model"] == "mock-flat"

    two = {"analysis": {"provider": "mock", "model": "mock-a"},
           "text": {"provider": "mock", "model": "mock-t"}}
    with open(storage.SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(two, f, ensure_ascii=False)
    loaded2 = storage.load_llm_settings()
    assert loaded2["analysis"]["model"] == "mock-a"
    assert loaded2["text"]["model"] == "mock-t"
    print("OK llm settings backward compat")


def test_methodology_excerpts(tmp_path, monkeypatch):
    """RAG по Методике: выдержки находятся; отсутствие файла → "" (fallback)."""
    from tm_core import methodology

    excerpts = methodology.excerpts_for("негативные последствия ущерб")
    assert excerpts.strip(), "По запросу должны находиться выдержки из Методики"
    assert len(excerpts) <= 4500, "Выдержки должны обрезаться по max_chars"

    monkeypatch.setattr(methodology, "METHODOLOGY_FILE", tmp_path / "no_such_file.txt")
    assert methodology.load_chunks() == [], "Отсутствующий файл → пустой список фрагментов"
    assert methodology.excerpts_for("негативные последствия ущерб") == "", \
        "Отсутствующий файл → пустые выдержки (fallback)"
    print(f"OK methodology: выдержек {len(excerpts)} символов, fallback работает")


def test_two_llm_configs():
    """Задание с РАЗНЫМИ конфигурациями для analysis и text: JSON-задачи
    (УБИ/тактики) уходят в analysis-провайдер, тексты разделов и сценарии —
    в text-провайдер."""
    profile = make_test_profile()
    state = storage.new_job(profile, "тест двух конфигураций LLM")
    suf = uuid.uuid4().hex[:8]  # уникальные имена моделей — мимо кэша LLM
    cfg_analysis = {"provider": "splitmock", "model": f"split-analysis-{suf}"}
    cfg_text = {"provider": "splitmock", "model": f"split-text-{suf}"}

    job = start_or_resume(state.job_id, cfg_analysis, cfg_text, ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"
    assert job.provider is not None and job.text_provider is not None
    assert job.provider.cfg.model == cfg_analysis["model"]
    assert job.text_provider.cfg.model == cfg_text["model"]

    a = SplitMockProvider.calls.get(cfg_analysis["model"], {"json": 0, "text": 0})
    t = SplitMockProvider.calls.get(cfg_text["model"], {"json": 0, "text": 0})
    assert a["json"] > 0, "JSON-задачи (УБИ/тактики) должны идти в analysis-провайдер"
    assert a["text"] == 0, "У analysis-провайдера не должно быть текстовых вызовов"
    assert t["text"] > 0, "Тексты разделов/сценариев должны идти в text-провайдер"
    assert t["json"] == 0, "У text-провайдера не должно быть JSON-вызовов"
    print(f"OK two llm configs: analysis json={a['json']}, text texts={t['text']}")
    return final


def test_security_tools_dict():
    """Справочник СЗИ/СКЗИ: читается, ≥ 20 средств, is_crypto_tool определяет
    crypto-средства по строке выбора UI."""
    from tm_core import dictionaries as dicts

    data = dicts.security_tools()
    assert data.get("categories"), "В справочнике должны быть категории"
    choices = dicts.security_tools_choices()
    assert len(choices) >= 20, f"Ожидалось ≥ 20 средств, получено {len(choices)}"
    assert all(isinstance(c, str) and c.strip() for c in choices)

    crypto_choice = next(c for c in choices if c.startswith("КриптоПро CSP"))
    assert dicts.is_crypto_tool(crypto_choice) is True, \
        f"«{crypto_choice}» — средство криптографической защиты (crypto=true)"
    plain_choice = next(c for c in choices if c.startswith("Dallas Lock"))
    assert dicts.is_crypto_tool(plain_choice) is False
    assert dicts.is_crypto_tool("Несуществующее средство") is False
    assert dicts.is_crypto_tool("") is False
    print(f"OK security_tools dict: {len(choices)} средств, crypto определяется")


def test_profile_security_tools_compat():
    """Обратная совместимость: старая строка security_tools разбивается по «;»;
    to_dict/from_dict roundtrip; has_crypto по умолчанию False."""
    p = Profile.from_dict({"object_name": "Тест", "security_tools": "Антивирус; МЭ"})
    assert p.security_tools == ["Антивирус", "МЭ"], \
        f"Строка должна разбиваться в список: {p.security_tools}"
    assert p.has_crypto is False

    p2 = Profile(object_name="Тест",
                 security_tools=["КриптоПро CSP (Крипто-Про)", "Secret Net Studio"],
                 has_crypto=True)
    d = p2.to_dict()
    assert d["security_tools"] == p2.security_tools and d["has_crypto"] is True
    p3 = Profile.from_dict(d)
    assert p3.security_tools == p2.security_tools and p3.has_crypto is True, \
        "Roundtrip to_dict/from_dict должен сохранять список и has_crypto"
    assert "КриптоПро CSP (Крипто-Про); Secret Net Studio" in p3.summary()

    p4 = Profile.from_dict({"security_tools": None})
    assert p4.security_tools == []
    print("OK profile security_tools compat")


def test_skzi_assessment():
    """Приказ ФСБ № 378: таблица «УЗ → типы угроз → класс СКЗИ»."""
    from tm_core.crypto_logic import skzi_assessment

    r = skzi_assessment(Profile(protection_level="УЗ-2", has_crypto=True))
    assert r is not None and r["uz"] == "УЗ-2"
    assert r["threat_types"] == [1, 2, 3]
    assert r["rows"] == [{"type": 1, "class": "КА"},
                         {"type": 2, "class": "КВ и выше"},
                         {"type": 3, "class": "КС1 и выше"}]
    assert r["min_class"] == "КА", "Наиболее строгий из требуемых — КА"

    r4 = skzi_assessment(Profile(protection_level="УЗ-4", has_crypto=True))
    assert r4 is not None and r4["threat_types"] == [3]
    assert r4["min_class"] == "КС1 и выше"

    assert skzi_assessment(Profile(protection_level="УЗ-2")) is None, \
        "Без has_crypto блок неприменим"
    assert skzi_assessment(Profile(protection_level="К1", has_crypto=True)) is None, \
        "Для класса ГИС (не УЗ-) блок неприменим"
    print("OK skzi_assessment")


def test_appendix5_and_73():
    """Полный прогон (mock) ИСПДн УЗ-2 + СКЗИ: в DOCX — подраздел 7.3,
    приказ № 378 в разделе 1 и Приложение 5 только с актуальными УБИ."""
    from docx import Document

    profile = make_test_profile()
    profile.protection_level = "УЗ-2"
    profile.has_crypto = True
    profile.security_tools = ["КриптоПро CSP (Крипто-Про)", "Антивирус"]
    state = storage.new_job(profile, "тест приложения 5 и 7.3")

    job = start_or_resume(state.job_id, {"provider": "mock", "model": "mock-app5"},
                          ubi_full_scan=False)
    final = job.run()
    assert final.status == "done", f"Статус: {final.status}, ошибка: {final.error}"

    s7 = storage.load_stage_result(state.job_id, "s7_ubi")
    assert s7.get("skzi"), "Результат s7 должен содержать блок skzi (приказ 378)"
    assert s7["skzi"]["min_class"] == "КА"

    doc = Document(final.docx_path)
    full_text = "\n".join(par.text for par in doc.paragraphs)
    assert "Приложение 5" in full_text, "Нет заголовка Приложения 5"
    assert "Перечень актуальных угроз безопасности информации" in full_text
    assert "378" in full_text, "Приказ ФСБ № 378 должен упоминаться (раздел 1 / 7.3)"
    assert "7.3. Тип актуальных угроз и требуемый класс СКЗИ" in full_text
    assert "КриптоПро CSP" in full_text, "СЗИ из списка профиля должны попасть в 2.3"

    # Приложение 5 — последняя таблица документа, только актуальные УБИ
    app5 = doc.tables[-1]
    codes_in_app5 = [row.cells[0].text.strip() for row in app5.rows][1:]  # без шапки
    actual_codes = sorted(i.get("ubi_code", f"УБИ.{int(i['Number']):03d}")
                          for i in s7["actual"])
    assert codes_in_app5 == actual_codes, \
        f"Приложение 5 должно содержать в точности актуальные УБИ: " \
        f"{codes_in_app5} != {actual_codes}"
    # конкретные номера — из результата s7 (mock: актуальны номера НЕ кратные 3)
    actual_nums = {int(i["Number"]) for i in s7["actual"] if str(i["Number"]).isdigit()}
    assert actual_nums, "В тесте должны быть актуальные УБИ"
    assert all(n % 3 != 0 for n in actual_nums), \
        "Mock-провайдер: актуальны только номера, не кратные 3"
    first_code = f"УБИ.{min(actual_nums):03d}"
    assert first_code in codes_in_app5, \
        f"{first_code} (mock: актуальна) должна быть в перечне"
    # неактуальные УБИ (в т.ч. кратные 3, если дошли до оценки) — отсутствуют
    not_actual = [i for i in s7["all"] if not i.get("matches")]
    mult3 = [i for i in not_actual
             if str(i["Number"]).isdigit() and int(i["Number"]) % 3 == 0]
    if mult3:  # например, «УБИ.003», если она рассматривалась
        code3 = mult3[0].get("ubi_code", f"УБИ.{int(mult3[0]['Number']):03d}")
        assert code3 not in codes_in_app5, \
            f"{code3} (номер кратен 3 — неактуальна) не должна попасть в перечень"
    for i in not_actual:
        code = i.get("ubi_code", f"УБИ.{int(i['Number']):03d}")
        assert code not in codes_in_app5, f"{code} неактуальна, но попала в Приложение 5"
    print(f"OK appendix5+7.3: актуальных УБИ в Приложении 5 — {len(codes_in_app5)}, "
          f"min_class={s7['skzi']['min_class']}")
    return final


def test_sync_interface_flags():
    """sync_interface_flags: отмеченные интерфейсы включают флаги (OR-логика),
    ручные включения не сбрасываются."""
    from tm_core.schema import sync_interface_flags

    p = Profile(network_interfaces=["Интернет", "Wi-Fi/беспроводные сети"])
    sync_interface_flags(p)
    assert p.has_internet is True
    assert p.has_wireless is True
    assert p.has_remote_access is False, "Не отмеченный интерфейс не должен включать флаг"
    assert p.has_external_integrations is False

    p2 = Profile(has_remote_access=True, network_interfaces=[])
    sync_interface_flags(p2)
    assert p2.has_remote_access is True, "Ручное включение флага должно сохраняться"
    assert p2.has_internet is False

    # выбор «Удаленный доступ» и «API» включает соответствующие флаги
    p3 = Profile(network_interfaces=["Удаленный доступ", "API (внешние интеграции)",
                                     "ЛВС"])
    sync_interface_flags(p3)
    assert p3.has_remote_access is True
    assert p3.has_external_integrations is True
    assert p3.has_internet is False
    print("OK sync_interface_flags")


def test_sample_profiles(tmp_path, monkeypatch):
    """Примеры профилей копируются в PROFILES_DIR при ensure_dirs, валидно
    загружаются, повторный вызов не затирает изменённый файл."""
    monkeypatch.setattr(storage, "OUTPUT_DIR", tmp_path / "output")
    monkeypatch.setattr(storage, "PROFILES_DIR", tmp_path / "output" / "profiles")
    monkeypatch.setattr(storage, "JOBS_DIR", tmp_path / "output" / "jobs")

    storage.ensure_dirs()
    copied = sorted(storage.PROFILES_DIR.glob("*.json"))
    assert len(copied) == 4, f"Должно быть скопировано 4 примера, есть: {[p.name for p in copied]}"

    for path in copied:
        with open(path, encoding="utf-8") as f:
            profile = Profile.from_dict(json.load(f))
        errors = profile.validate()  # не должно кидать исключений
        assert isinstance(errors, list)
        assert profile.object_name, f"{path.name}: нет наименования объекта"

    # повторный вызов не затирает изменённый файл
    target = copied[0]
    marker = {"object_name": "ИЗМЕНЕНО ПОЛЬЗОВАТЕЛЕМ"}
    with open(target, "w", encoding="utf-8") as f:
        json.dump(marker, f, ensure_ascii=False)
    storage.ensure_dirs()
    with open(target, encoding="utf-8") as f:
        assert json.load(f) == marker, "Изменённый пользователем профиль не должен затираться"
    print(f"OK sample profiles: {[p.name for p in copied]}")


if __name__ == "__main__":
    test_safe_parse_json_repairs()
    test_norm_ubi_number()
    test_full_pipeline()
    test_resume()
    test_retry_recovers()
    test_total_failure_marked()
    test_resume_reprocesses_failed()
    test_s8_old_format_partials_reprocessed()
    test_downstream_invalidation()
    test_draft_roundtrip()
    test_llm_settings_backward_compat()
    test_two_llm_configs()
    print("Все тесты пройдены.")
